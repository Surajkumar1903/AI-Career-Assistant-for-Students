"""
Export tailored job-matched resume to PDF (ReportLab) or DOCX (python-docx).
Classic single-page CV layout: centered name + contact bar, section titles + rules, black body text.
"""
from __future__ import annotations

import os
import re
from datetime import datetime
from xml.sax.saxutils import escape


def _para_html(text: str) -> str:
    if not text or not str(text).strip():
        return "—"
    e = escape(str(text).strip())
    return e.replace("\n", "<br/>")


def _links_from_resume_text(text: str) -> dict[str, str]:
    """Pull obvious profile URLs from resume excerpt (no fabrication)."""
    out: dict[str, str] = {}
    if not text:
        return out
    t = text.replace("\n", " ")
    m = re.search(r"((?:https?://)?(?:www\.)?github\.com/[\w./-]+)", t, re.I)
    if m:
        url = m.group(1)
        if not url.startswith("http"):
            url = "https://" + url.lstrip("/")
        out["github"] = url.rstrip(").,;]")
    m = re.search(r"((?:https?://)?(?:www\.)?linkedin\.com/in/[\w%-]+)", t, re.I)
    if m:
        url = m.group(1)
        if not url.startswith("http"):
            url = "https://" + url.lstrip("/")
        out["linkedin"] = url.rstrip(").,;]")
    m = re.search(r"(https?://[\w.-]*(?:portfolio|behance|dribbble)[\w./-]*)", t, re.I)
    if m:
        out["portfolio"] = m.group(1).rstrip(").,;]")
    return out


def _contact_line_parts(user, bundle: dict) -> list[str]:
    rc = bundle.get("resume_contact") or {}
    raw = bundle.get("resume_raw_excerpt") or ""
    links = _links_from_resume_text(raw)

    phone = (getattr(user, "phone", None) or rc.get("phone") or "").strip()
    email = (getattr(user, "email", None) or rc.get("email") or "").strip()

    parts: list[str] = []
    if phone:
        parts.append(phone)
    if email:
        parts.append(email)
    if links.get("portfolio"):
        parts.append(links["portfolio"])
    if links.get("github"):
        parts.append(links["github"])
    if links.get("linkedin"):
        parts.append(links["linkedin"])
    return parts


def _section_body_paragraphs(val: str | list | None, key: str) -> list[str]:
    if key == "recommended_skills_to_learn" and isinstance(val, list):
        return [f"• {x}" for x in val if str(x).strip()] or ["—"]
    s = str(val or "").strip()
    if not s:
        return ["—"]
    lines = [ln.strip() for ln in s.split("\n") if ln.strip()]
    return lines if lines else [s]


def build_job_resume_pdf(bundle: dict, user, output_path: str) -> str:
    """Single-page style CV: black typography, rules under section titles, centered header."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        HRFlowable,
        Table,
        TableStyle,
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    sections = bundle.get("sections") or {}
    job_role = bundle.get("job_role") or ""
    match = bundle.get("match") or {}
    ats = int(match.get("ats_match_score", 0) or 0)

    sp = bundle.get("style_profile") or {}
    name_pt = float(sp.get("name_pt", 20))
    contact_pt = float(sp.get("contact_pt", 9))
    section_pt = float(sp.get("section_pt", 11))
    body_pt = float(sp.get("body_pt", 10))
    m_cm = float(sp.get("margins_cm", 1.45))
    line_lead = float(sp.get("line_leading", body_pt * 1.35))

    total_chars = sum(len(str(sections.get(k, "") or "")) for k in sections)
    if sp.get("one_page_compact") and total_chars > 7200:
        body_pt = max(8.5, round(body_pt * 0.92, 1))
        line_lead = round(body_pt * 1.32, 1)

    BLACK = colors.black
    GRAY = colors.HexColor("#333333")
    RULE = colors.HexColor("#000000")

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=m_cm * cm,
        leftMargin=m_cm * cm,
        topMargin=max(1.0, m_cm * 0.85) * cm,
        bottomMargin=max(1.0, m_cm * 0.85) * cm,
    )

    name_style = ParagraphStyle(
        "CVName",
        fontName="Helvetica-Bold",
        fontSize=name_pt,
        alignment=TA_CENTER,
        textColor=BLACK,
        spaceAfter=6,
        leading=name_pt * 1.15,
    )
    contact_style = ParagraphStyle(
        "CVContact",
        fontName="Helvetica",
        fontSize=contact_pt,
        alignment=TA_CENTER,
        textColor=GRAY,
        spaceAfter=12,
        leading=contact_pt * 1.2,
    )
    sec_title = ParagraphStyle(
        "CVSec",
        fontName="Helvetica-Bold",
        fontSize=section_pt,
        alignment=TA_LEFT,
        textColor=BLACK,
        spaceBefore=6,
        spaceAfter=2,
        leading=section_pt * 1.25,
    )
    body = ParagraphStyle(
        "CVBody",
        fontName="Helvetica",
        fontSize=body_pt,
        alignment=TA_LEFT,
        textColor=BLACK,
        leading=line_lead,
        spaceAfter=2,
    )
    body_justify = ParagraphStyle(
        "CVBodyJust",
        parent=body,
        alignment=TA_JUSTIFY,
    )
    foot = ParagraphStyle(
        "CVFoot",
        fontName="Helvetica",
        fontSize=7.2,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#666666"),
        spaceBefore=8,
    )

    story: list = []
    display_name = (user.full_name or user.username or "Candidate").strip().upper()
    story.append(Paragraph(escape(display_name), name_style))

    c_parts = _contact_line_parts(user, bundle)
    if c_parts:
        contact_html = " &nbsp;|&nbsp; ".join(escape(p) for p in c_parts)
        story.append(Paragraph(contact_html, contact_style))
    else:
        story.append(Spacer(1, 0.15 * cm))

    order = [
        ("Professional Summary", "professional_summary", True),
        ("Professional Experience", "experience", False),
        ("Projects", "projects", False),
        ("Technical Skills", "skills", False),
        ("Education", "education", False),
        ("Certifications", "certifications", False),
        ("Achievements", "achievements", False),
        ("Recommended Skills to Learn", "recommended_skills_to_learn", False),
    ]

    for title, key, justify in order:
        story.append(Paragraph(escape(title), sec_title))
        story.append(HRFlowable(width="100%", thickness=0.75, color=RULE, spaceBefore=0, spaceAfter=6))
        val = sections.get(key)
        if key == "recommended_skills_to_learn" and isinstance(val, list):
            lines = _section_body_paragraphs(val, key)
            for line in lines:
                story.append(Paragraph(escape(line), body))
            continue

        if key == "certifications":
            if isinstance(val, list):
                bullets = [f"• {str(x).strip()}" for x in val if str(x).strip()] or ["—"]
            else:
                raw_lines = [ln.strip() for ln in str(val or "").split("\n") if ln.strip()]
                bullets = []
                for ln in raw_lines:
                    bullets.append(ln if ln.startswith("•") else f"• {ln}")
                if not bullets:
                    bullets = ["—"]
            mid = (len(bullets) + 1) // 2
            col_a = bullets[:mid]
            col_b = bullets[mid:]
            data = []
            for i in range(max(len(col_a), len(col_b))):
                left = col_a[i] if i < len(col_a) else ""
                right = col_b[i] if i < len(col_b) else ""
                data.append([Paragraph(escape(left), body), Paragraph(escape(right), body)])
            tw = doc.width
            tbl = Table(data, colWidths=[tw * 0.48, tw * 0.48])
            tbl.setStyle(
                TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ]
                )
            )
            story.append(tbl)
            story.append(Spacer(1, 0.12 * cm))
            continue

        text = str(val if isinstance(val, str) else "").strip()
        if not text:
            story.append(Paragraph("—", body))
            continue
        style = body_justify if justify else body
        story.append(Paragraph(_para_html(text), style))

    story.append(Spacer(1, 0.15 * cm))
    ab = bundle.get("ats_before")
    aa = bundle.get("ats_after")
    if ab is not None and aa is not None:
        foot_line = (
            f"Target role: {job_role} · JD alignment estimate — Before: {int(ab)}/100 → After: {int(aa)}/100. "
            "Content from your upload only; recommended skills are learning gaps."
        )
    else:
        foot_line = (
            f"Target role: {job_role} · JD keyword alignment (estimate): {ats}/100 — "
            "Content sourced from your upload; recommended skills are learning gaps only."
        )
    story.append(Paragraph(escape(foot_line), foot))
    doc.build(story)
    return output_path


def _docx_set_bottom_border(paragraph):
    """Thin bottom rule under a paragraph (section title style)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_job_resume_docx(bundle: dict, user, output_path: str) -> str:
    """DOCX mirroring the PDF CV structure."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    sections = bundle.get("sections") or {}
    job_role = bundle.get("job_role") or ""
    match = bundle.get("match") or {}
    ats = int(match.get("ats_match_score", 0) or 0)

    sp = bundle.get("style_profile") or {}
    name_pt = float(sp.get("name_pt", 20))
    contact_pt = float(sp.get("contact_pt", 9))
    section_pt = float(sp.get("section_pt", 11))
    body_pt = float(sp.get("body_pt", 10))
    margin_pt = max(36, min(52, int(float(sp.get("margins_cm", 1.45)) * 28.346)))  # cm to pt ~28.35 pt/cm

    d = Document()
    sec = d.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Pt(margin_pt)

    # Name
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run((user.full_name or user.username or "Candidate").strip().upper())
    r.bold = True
    r.font.size = Pt(name_pt)
    r.font.color.rgb = RGBColor(0, 0, 0)

    c_parts = _contact_line_parts(user, bundle)
    if c_parts:
        pc = d.add_paragraph(" | ".join(c_parts))
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in pc.runs:
            run.font.size = Pt(contact_pt)
            run.font.color.rgb = RGBColor(51, 51, 51)
    d.add_paragraph()

    order = [
        ("Professional Summary", "professional_summary", True),
        ("Professional Experience", "experience", False),
        ("Projects", "projects", False),
        ("Technical Skills", "skills", False),
        ("Education", "education", False),
        ("Certifications", "certifications", False),
        ("Achievements", "achievements", False),
        ("Recommended Skills to Learn", "recommended_skills_to_learn", False),
    ]

    for title, key, justify in order:
        ph = d.add_paragraph()
        rh = ph.add_run(title)
        rh.bold = True
        rh.font.size = Pt(section_pt)
        _docx_set_bottom_border(ph)
        ph.paragraph_format.space_after = Pt(8)

        val = sections.get(key)
        if key == "recommended_skills_to_learn" and isinstance(val, list):
            for x in val or []:
                d.add_paragraph(str(x), style="List Bullet")
            if not val:
                d.add_paragraph("—")
            continue

        if key == "certifications" and isinstance(val, str):
            lines = [ln.strip() for ln in val.split("\n") if ln.strip()]
            bullets = [ln if ln.startswith("•") else f"• {ln}" for ln in lines] or ["—"]
            mid = (len(bullets) + 1) // 2
            left = bullets[:mid]
            right = bullets[mid:]
            n = max(len(left), len(right), 1)
            tbl = d.add_table(rows=n, cols=2)
            for i in range(n):
                tbl.rows[i].cells[0].text = left[i] if i < len(left) else ""
                tbl.rows[i].cells[1].text = right[i] if i < len(right) else ""
            continue

        if key == "certifications" and isinstance(val, list):
            bullets = [f"• {str(x).strip()}" for x in val if str(x).strip()] or ["—"]
            mid = (len(bullets) + 1) // 2
            left = bullets[:mid]
            right = bullets[mid:]
            n = max(len(left), len(right), 1)
            tbl = d.add_table(rows=n, cols=2)
            for i in range(n):
                tbl.rows[i].cells[0].text = left[i] if i < len(left) else ""
                tbl.rows[i].cells[1].text = right[i] if i < len(right) else ""
            continue

        text = str(val or "").strip()
        if not text:
            d.add_paragraph("—")
            continue
        for block in text.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            pb = d.add_paragraph(block)
            pb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
            for run in pb.runs:
                run.font.size = Pt(body_pt)
                run.font.name = "Calibri"

    fn = d.add_paragraph()
    fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ab = bundle.get("ats_before")
    aa = bundle.get("ats_after")
    if ab is not None and aa is not None:
        foot_txt = (
            f"Target role: {job_role} · JD alignment estimate — Before: {int(ab)}/100 → After: {int(aa)}/100. "
            "Sourced from your upload; recommended skills are gaps only."
        )
    else:
        foot_txt = (
            f"Target role: {job_role} · JD alignment (estimate): {ats}/100 — "
            "Sourced from your upload; recommended skills are gaps only."
        )
    rr = fn.add_run(foot_txt)
    rr.italic = True
    rr.font.size = Pt(8)
    rr.font.color.rgb = RGBColor(102, 102, 102)

    d.save(output_path)
    return output_path
